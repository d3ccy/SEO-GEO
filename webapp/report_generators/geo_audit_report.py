"""
Parameterized GEO Audit Report generator.
Produces a full branded DOCX GEO audit report for any client/URL.

params dict keys (all optional with sensible defaults):
  client_name   - e.g. "Numiko"
  client_domain - e.g. "numiko.com"
  project_name  - e.g. "Website Relaunch 2026"
  date          - e.g. "February 2026"
  logo_path     - absolute path to logo image (falls back to agency logo)

audit dict keys (from audit_service.run_audit()):
  url, title, title_length, title_ok,
  description, description_length, description_ok,
  og_tags, h1, jsonld_count,
  load_time, load_time_ok,
  robots_exists, ai_bots, ai_bots_blocked, has_sitemap, sitemap_url, score,
  backlinks_rank, referring_domains, total_backlinks
"""
import os
import logging
from datetime import datetime
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx import Document

from .docx_helpers import (
    BLACK, DARK, GRAY, LIGHT_GRAY, WHITE, RED, GREEN, AMBER,
    add_styled_para, add_heading, add_bullet, add_table,
    add_callout_box, set_cell_shading, create_document,
)

logger = logging.getLogger(__name__)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DEFAULT_LOGO = os.path.join(BASE_DIR, 'fonts', 'numiko_logo.png')


def _score_color(score: int) -> str:
    if score >= 71:
        return '167832'
    if score >= 41:
        return 'C87814'
    return 'B42828'


def _status_color(ok: bool, warn: bool = False):
    if ok:
        return GREEN
    if warn:
        return AMBER
    return RED


def _yes_no(val) -> str:
    return 'Yes' if val else 'No'


def _geo_method_present(score: int) -> str:
    """Map raw audit data to a rough content-method count."""
    if score >= 71:
        return '6–9'
    if score >= 41:
        return '3–5'
    return '0–2'


def build_geo_audit_report(params: dict, audit: dict) -> Document:
    client_name = params.get('client_name') or 'Client'
    client_domain = params.get('client_domain') or audit.get('url', '')
    project_name = params.get('project_name') or ''
    date_str = params.get('date') or datetime.now().strftime('%B %Y')
    logo_path = params.get('logo_path') or DEFAULT_LOGO

    # Convenience aliases
    score = audit.get('score', 0)
    title = audit.get('title') or ''
    description = audit.get('description') or ''
    h1 = audit.get('h1') or ''
    og_tags = audit.get('og_tags', False)
    jsonld_count = audit.get('jsonld_count', 0)
    load_time = audit.get('load_time')
    load_time_ok = audit.get('load_time_ok', False)
    robots_exists = audit.get('robots_exists', False)
    ai_bots = audit.get('ai_bots', [])
    ai_bots_blocked = audit.get('ai_bots_blocked', [])
    has_sitemap = audit.get('has_sitemap', False)
    sitemap_url = audit.get('sitemap_url') or '/sitemap.xml'
    url = audit.get('url', client_domain)
    backlinks_rank = audit.get('backlinks_rank')
    referring_domains = audit.get('referring_domains')
    total_backlinks = audit.get('total_backlinks')

    score_hex = _score_color(score)
    score_rgb = (int(score_hex[0:2], 16), int(score_hex[2:4], 16), int(score_hex[4:6], 16))
    from docx.shared import RGBColor
    SCORE_COLOR = RGBColor(*score_rgb)

    doc = create_document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, height=Cm(3))

    doc.add_paragraph()
    add_styled_para(doc, 'GEO Audit Report', size=28, bold=True, color=BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_styled_para(doc, client_domain, size=18, bold=True, color=BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_styled_para(doc, client_name, size=14, color=GRAY,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    if project_name:
        add_styled_para(doc, project_name, size=11, bold=True, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_styled_para(doc, date_str, size=10, color=LIGHT_GRAY,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_styled_para(doc, 'Generative Engine Optimization', size=10, color=LIGHT_GRAY,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))
    data_sources = 'Data sources: Direct site crawl, robots.txt, sitemap.xml'
    if backlinks_rank is not None or referring_domains is not None:
        data_sources += ', DataForSEO Backlinks API'
    add_styled_para(doc, data_sources,
                    size=8, color=LIGHT_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=Pt(2))
    add_styled_para(doc, 'Based on Princeton GEO research framework (KDD 2024)',
                    size=8, color=LIGHT_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ───────────────────────────────────────────────────
    add_heading(doc, 'Executive Summary', level=1)

    if score >= 71:
        summary_rating = 'a strong'
        summary_outlook = (
            'The site has a solid GEO foundation. Key strengths include good meta tag '
            'coverage, structured data, and fast page load times. The priority is '
            'enriching content with the Princeton GEO methods — statistics, expert quotes, '
            'citations, and FAQ formatting — to maximise AI citation rates.'
        )
    elif score >= 41:
        summary_rating = 'a moderate'
        summary_outlook = (
            'The site has some GEO foundations in place but significant gaps remain. '
            'Priority areas include schema markup, AI bot access, and content enrichment '
            'with the Princeton GEO methods to improve AI citation visibility.'
        )
    else:
        summary_rating = 'a weak'
        summary_outlook = (
            'The site has significant GEO gaps that limit its visibility in AI-generated '
            'answers. Immediate action is needed on technical fundamentals — schema markup, '
            'meta tags, AI bot access — before content optimisation can be effective.'
        )

    add_styled_para(doc, (
        f'{client_name} ({client_domain}) scores {score}/100 for GEO readiness — '
        f'{summary_rating} foundation for AI search visibility. {summary_outlook}'
    ))

    add_styled_para(doc, (
        'This audit is based on the Princeton University GEO research framework '
        '(arXiv:2311.09735, KDD 2024), which identified 9 content optimisation methods '
        'that improve AI visibility by up to 40%. GEO (Generative Engine Optimization) is '
        'the practice of optimising website content to be cited by AI search engines — '
        'ChatGPT, Perplexity, Google AI Overview, Microsoft Copilot, and Claude.'
    ))

    # Score box
    p = doc.add_paragraph()
    run = p.add_run(f'GEO Score: {score} / 100')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = SCORE_COLOR
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)

    # Key findings table
    add_heading(doc, 'Key Findings at a Glance', level=2)

    findings = []
    findings.append(['Page Title', title[:60] + ('…' if len(title) > 60 else '') if title else 'Missing',
                     'Good' if (title and audit.get('title_ok')) else ('Weak' if title else 'Missing')])
    findings.append(['Meta Description',
                     description[:60] + ('…' if len(description) > 60 else '') if description else 'Missing',
                     'Good' if (description and audit.get('description_ok')) else ('Weak' if description else 'Missing')])
    findings.append(['H1 Heading', h1[:50] + ('…' if len(h1) > 50 else '') if h1 else 'Missing',
                     'Good' if h1 else 'Missing'])
    findings.append(['OG / Social Tags', 'Present' if og_tags else 'Missing',
                     'Good' if og_tags else 'Missing'])
    findings.append(['JSON-LD Schema', f'{jsonld_count} block{"s" if jsonld_count != 1 else ""}',
                     'Good' if jsonld_count > 0 else 'Missing'])
    findings.append(['Page Load Time', f'{load_time}s' if load_time else 'Unknown',
                     'Good' if load_time_ok else ('Weak' if load_time else 'Missing')])
    findings.append(['robots.txt', 'Found' if robots_exists else 'Not found',
                     'Good' if robots_exists else 'Missing'])
    # AI bots: show allowed and blocked
    if ai_bots and ai_bots_blocked:
        ai_bots_value = f'Allowed: {", ".join(ai_bots)}; Blocked: {", ".join(ai_bots_blocked)}'
        ai_bots_status = 'Partial'
    elif ai_bots:
        ai_bots_value = ', '.join(ai_bots)
        ai_bots_status = 'Good'
    elif ai_bots_blocked:
        ai_bots_value = f'Blocked: {", ".join(ai_bots_blocked)}'
        ai_bots_status = 'Warning'
    else:
        ai_bots_value = 'Not configured'
        ai_bots_status = 'Missing'
    findings.append(['AI Bot Access', ai_bots_value, ai_bots_status])
    findings.append(['XML Sitemap', 'Found' if has_sitemap else 'Not found',
                     'Good' if has_sitemap else 'Missing'])
    # Backlinks (if available)
    if backlinks_rank is not None:
        findings.append(['Domain Rank', str(backlinks_rank), 'Info'])
    if referring_domains is not None:
        findings.append(['Referring Domains', f'{referring_domains:,}', 'Info'])
    if total_backlinks is not None:
        findings.append(['Total Backlinks', f'{total_backlinks:,}', 'Info'])

    add_table(doc, ['Check', 'Value', 'Status'], findings)

    doc.add_page_break()

    # ── WHAT IS GEO ─────────────────────────────────────────────────────────
    add_heading(doc, 'What is GEO?', level=1)

    add_styled_para(doc, (
        'Generative Engine Optimization (GEO) is the practice of optimising website content to be '
        'cited by AI search engines — ChatGPT, Perplexity, Google AI Overview, Microsoft Copilot, '
        'and Claude. Unlike traditional SEO where the goal is to rank on a results page, GEO focuses '
        'on being the source that AI platforms quote when answering user questions.'
    ))

    add_styled_para(doc, (
        'Being cited is the new "ranking #1." When a user asks an AI engine a question relevant to '
        f'{client_name}\'s services, the AI synthesises an answer from multiple sources. Sites that '
        'provide extractable facts — statistics, expert quotes, structured data, citations — are the '
        'ones that get cited. Marketing copy, no matter how well-written, is rarely cited directly.'
    ))

    add_styled_para(doc, (
        'This audit is based on the Princeton University GEO research framework (arXiv:2311.09735, '
        'accepted KDD 2024), which identified 9 methods that improve AI visibility by up to 40%.'
    ))

    add_table(doc,
              ['Method', 'AI Visibility Boost', 'Description'],
              [
                  ['Cite Sources', '+40%', 'Add authoritative citations and references'],
                  ['Statistics Addition', '+37%', 'Include specific numbers and data points'],
                  ['Quotation Addition', '+30%', 'Add expert quotes with attribution'],
                  ['Authoritative Tone', '+25%', 'Use confident, expert language'],
                  ['Easy-to-understand', '+20%', 'Simplify complex concepts'],
                  ['Technical Terms', '+18%', 'Include domain-specific terminology'],
                  ['Unique Words', '+15%', 'Increase vocabulary diversity'],
                  ['Fluency Optimization', '+15–30%', 'Improve readability and flow'],
                  ['Keyword Stuffing', '−10%', 'AVOID — hurts AI visibility'],
              ])

    add_styled_para(doc, 'Best combination: Fluency + Statistics = Maximum boost',
                    bold=True, color=BLACK)

    doc.add_page_break()

    # ── 1. META TAGS & ON-PAGE FUNDAMENTALS ─────────────────────────────────
    add_heading(doc, '1. Meta Tags & On-Page Fundamentals', level=1)

    add_styled_para(doc, (
        'Meta tags are the first signals AI crawlers use to understand a page. '
        'A missing or poor title, description, or H1 directly reduces the chance of '
        'AI engines citing your content correctly.'
    ))

    # Title
    add_heading(doc, 'Page Title', level=2)
    if title:
        add_styled_para(doc, f'Found: "{title}"')
        add_styled_para(doc, f'Length: {audit.get("title_length", len(title))} characters '
                        f'({"good — under 60" if audit.get("title_ok") else "too long — aim for under 60 characters"})')
        if not audit.get('title_ok'):
            add_callout_box(doc, 'Recommendation', [
                'Shorten the page title to under 60 characters. '
                'AI engines truncate long titles and may misrepresent the page topic. '
                'Lead with the most important keyword or brand name.'
            ], 'C87814', 'FFF8E6')
    else:
        add_styled_para(doc, 'No page title found.', bold=True, color=RED)
        add_callout_box(doc, 'Critical: Missing Page Title', [
            'Every page must have a unique, descriptive <title> tag. '
            'AI engines use the title as the primary label for the page. '
            'A missing title means AI platforms cannot reliably identify or cite this page.'
        ], 'B42828', 'FFF0F0')

    # Description
    add_heading(doc, 'Meta Description', level=2)
    if description:
        add_styled_para(doc, f'Found: "{description[:120]}{"…" if len(description) > 120 else ""}"')
        add_styled_para(doc, f'Length: {audit.get("description_length", len(description))} characters '
                        f'({"good — under 155" if audit.get("description_ok") else "too long — aim for under 155 characters"})')
    else:
        add_styled_para(doc, 'No meta description found.', bold=True, color=RED)
        add_callout_box(doc, 'Recommendation: Add Meta Description', [
            'Write a compelling 120–155 character meta description for every page. '
            'While Google does not always use the meta description in SERPs, AI engines '
            'use it as a summary signal when deciding whether to cite the page. '
            'Include your primary keyword and a clear value proposition.'
        ], 'C87814', 'FFF8E6')

    # H1
    add_heading(doc, 'H1 Heading', level=2)
    if h1:
        add_styled_para(doc, f'Found: "{h1}"')
        add_styled_para(doc, 'The H1 heading is present — good. Ensure it is descriptive and '
                        'matches the page\'s primary topic.')
    else:
        add_styled_para(doc, 'No H1 heading found.', bold=True, color=RED)
        add_callout_box(doc, 'Recommendation: Add an H1 Heading', [
            'Every page should have exactly one H1 heading that clearly states the page topic. '
            'AI engines treat the H1 as a strong relevance signal — it anchors the semantic '
            'context of the entire page. Missing H1s reduce citation accuracy.'
        ], 'B42828', 'FFF0F0')

    # OG Tags
    add_heading(doc, 'Open Graph (Social / OG) Tags', level=2)
    if og_tags:
        add_styled_para(doc, 'Open Graph tags are present — good.')
        add_styled_para(doc, (
            'OG tags help AI platforms (and social media previews) understand the content. '
            'Ensure og:title, og:description, og:image, and og:url are set correctly on all pages.'
        ))
    else:
        add_styled_para(doc, 'No Open Graph tags found.', color=AMBER)
        add_callout_box(doc, 'Recommendation: Add OG Tags', [
            'Add Open Graph meta tags to improve how content is shared and interpreted. '
            'At minimum: og:title, og:description, og:url, og:image. '
            'These are used by social platforms, AI search engines, and link preview systems.'
        ], 'C87814', 'FFF8E6')

    doc.add_page_break()

    # ── 2. SCHEMA MARKUP ────────────────────────────────────────────────────
    add_heading(doc, '2. Schema Markup (JSON-LD)', level=1)

    add_styled_para(doc, (
        'JSON-LD structured data is one of the most powerful GEO signals. It tells AI engines '
        'exactly what type of entity the page represents, what it offers, and how to extract '
        'key facts. Sites without schema rely entirely on AI engines to infer context from '
        'prose — which is far less reliable.'
    ))

    add_heading(doc, 'Current State', level=2)

    if jsonld_count == 0:
        add_styled_para(doc, 'No JSON-LD schema blocks found on this page.', bold=True, color=RED)
        add_callout_box(doc, 'Critical: No Structured Data', [
            'Zero JSON-LD blocks were detected. This is one of the most impactful GEO gaps. '
            'AI engines use structured data to extract facts, understand entity types, and '
            'generate accurate answers. Without schema, all fact extraction depends on '
            'parsing unstructured prose — much less reliable and citable.',
        ], 'B42828', 'FFF0F0')
    elif jsonld_count == 1:
        add_styled_para(doc, f'{jsonld_count} JSON-LD block found.', color=AMBER)
        add_styled_para(doc, (
            'Some structured data is present, but a single generic block is unlikely to be '
            'sufficient. Review what is included and consider expanding to page-specific types.'
        ))
    else:
        add_styled_para(doc, f'{jsonld_count} JSON-LD blocks found — good.', color=GREEN)
        add_styled_para(doc, (
            'Multiple schema blocks are present. Ensure they are accurate, complete, '
            'and include page-specific types beyond just the organisation.'
        ))

    add_heading(doc, 'Recommended Schema Types', level=2)

    add_styled_para(doc, (
        'The following schema types are recommended based on the site type and '
        'the Princeton GEO research framework:'
    ))

    schema_rows = [
        ['Organization / LocalBusiness', 'All pages (sitewide)', 'High',
         'Name, address, phone, opening hours, social profiles, logo'],
        ['WebPage', 'All pages', 'High',
         'Name, description, breadcrumb, dateModified'],
        ['FAQPage', 'Key content pages', 'Critical',
         '+40% AI visibility boost — highest impact schema type'],
        ['BreadcrumbList', 'All pages', 'Medium',
         'Navigation structure for AI engines'],
        ['Article / BlogPosting', 'Blog / news', 'Medium',
         'datePublished, author, publisher for E-E-A-T'],
        ['Product / Service', 'Product / service pages', 'High',
         'Name, description, offers, review, aggregateRating'],
    ]

    add_table(doc, ['Schema Type', 'Where', 'Priority', 'Key Properties'], schema_rows)

    add_heading(doc, 'FAQPage Schema — Highest Impact', level=2)

    add_styled_para(doc, (
        'The Princeton GEO research identifies FAQ-structured content as the single highest-impact '
        'method, boosting AI visibility by up to +40%. FAQPage schema allows AI engines to extract '
        'specific question-and-answer pairs directly from your pages.'
    ))

    add_styled_para(doc, 'To implement FAQPage schema:', bold=True, color=BLACK)
    add_bullet(doc, 'Structure common questions on key pages using expandable accordion or Q&A sections')
    add_bullet(doc, 'Add FAQPage JSON-LD listing the questions and answers in structured format')
    add_bullet(doc, 'Target questions your audience would ask an AI: "What is X?", "How does Y work?", "How much does Z cost?"')
    add_bullet(doc, 'Each answer should be concise, factual, and directly answer the question')

    doc.add_page_break()

    # ── 3. AI BOT ACCESS ────────────────────────────────────────────────────
    add_heading(doc, '3. AI Bot Access', level=1)

    add_styled_para(doc, (
        'Before AI engines can cite your content, their crawlers must be able to access it. '
        'Two files control this: robots.txt (which bots are allowed) and sitemap.xml (which '
        'pages exist). Misconfigured bot access is a common but easily fixed GEO blocker.'
    ))

    # robots.txt
    add_heading(doc, 'robots.txt', level=2)

    if robots_exists:
        add_styled_para(doc, 'robots.txt found — good.')
    else:
        add_styled_para(doc, 'No robots.txt found.', bold=True, color=RED)
        add_callout_box(doc, 'Recommendation: Create robots.txt', [
            'A robots.txt file is expected by all web crawlers. Without one, '
            'crawlers may behave unpredictably. Create a robots.txt at the site root '
            'with explicit Allow directives for AI bots.'
        ], 'B42828', 'FFF0F0')

    add_heading(doc, 'AI Bot Directives', level=2)

    def _bot_status(bot):
        if bot in ai_bots:
            return 'Explicitly allowed'
        if bot in ai_bots_blocked:
            return 'BLOCKED — remove Disallow'
        if robots_exists:
            return 'Implicitly allowed (not mentioned)'
        return 'Unknown'

    bot_rows = [
        ['GPTBot', 'ChatGPT / OpenAI', 'Yes' if 'GPTBot' in ai_bots or 'GPTBot' in ai_bots_blocked else 'No', _bot_status('GPTBot')],
        ['ChatGPT-User', 'ChatGPT browsing', 'Yes' if 'ChatGPT-User' in ai_bots or 'ChatGPT-User' in ai_bots_blocked else 'No', _bot_status('ChatGPT-User')],
        ['ClaudeBot', 'Claude / Anthropic', 'Yes' if 'ClaudeBot' in ai_bots or 'ClaudeBot' in ai_bots_blocked else 'No', _bot_status('ClaudeBot')],
        ['anthropic-ai', 'Claude (legacy)', 'Yes' if 'anthropic-ai' in ai_bots or 'anthropic-ai' in ai_bots_blocked else 'No', _bot_status('anthropic-ai')],
        ['PerplexityBot', 'Perplexity', 'Yes' if 'PerplexityBot' in ai_bots or 'PerplexityBot' in ai_bots_blocked else 'No', _bot_status('PerplexityBot')],
        ['Googlebot', 'Google AI Overview', 'Standard', 'Implicitly allowed (standard)'],
        ['Bingbot', 'Copilot / Bing', 'Standard', 'Implicitly allowed (standard)'],
    ]

    add_table(doc, ['Bot', 'Platform', 'In robots.txt', 'Status'], bot_rows)

    if ai_bots_blocked:
        add_callout_box(doc, 'Critical: AI Bots Are Blocked', [
            f'The following AI crawlers are explicitly blocked in robots.txt: {", ".join(ai_bots_blocked)}.',
            'This prevents these AI platforms from indexing your content and citing it in answers.',
            '',
            'To fix: change the Disallow rule to Allow for each AI bot:',
            '',
            *[f'User-agent: {bot}\nAllow: /' for bot in ai_bots_blocked],
        ], 'B42828', 'FFF0F0')
    elif not ai_bots:
        add_callout_box(doc, 'Recommendation: Add Explicit AI Bot Directives', [
            'No AI-specific crawler rules were found in robots.txt. '
            'While AI bots are implicitly allowed if not blocked, explicitly allowing them '
            'sends a positive signal that the site welcomes AI crawler access.',
            '',
            'Add to robots.txt:',
            'User-agent: GPTBot',
            'Allow: /',
            '',
            'User-agent: ClaudeBot',
            'Allow: /',
            '',
            'User-agent: PerplexityBot',
            'Allow: /',
            '',
            'User-agent: ChatGPT-User',
            'Allow: /',
        ], 'C87814', 'FFF8E6')
    else:
        add_styled_para(doc, (
            f'AI bots explicitly allowed in robots.txt: {", ".join(ai_bots)}. '
            'This is a positive signal. Verify that no Crawl-delay is set '
            'that would slow AI indexing.'
        ))

    # Sitemap
    add_heading(doc, 'XML Sitemap', level=2)

    if has_sitemap:
        add_styled_para(doc, f'Sitemap found at {sitemap_url} — good.')
        add_styled_para(doc, (
            'Ensure the sitemap is up to date, includes all important pages, '
            'and is referenced in robots.txt (Sitemap: https://yourdomain.com/sitemap.xml).'
        ))
    else:
        add_styled_para(doc, 'No sitemap.xml found.', bold=True, color=RED)
        add_callout_box(doc, 'Recommendation: Create an XML Sitemap', [
            'An XML sitemap at /sitemap.xml tells AI crawlers which pages exist '
            'and when they were last updated. Without a sitemap, new or updated '
            'content may take much longer to be discovered and indexed.',
            '',
            'Generate a sitemap using your CMS plugin (Yoast for WordPress, '
            'simple_sitemap for Drupal, etc.) and submit it to Google Search Console.',
        ], 'B42828', 'FFF0F0')

    doc.add_page_break()

    # ── 4. PERFORMANCE ──────────────────────────────────────────────────────
    add_heading(doc, '4. Technical Performance', level=1)

    add_styled_para(doc, (
        'Page load time directly affects GEO readiness in two ways: slow sites are crawled '
        'less frequently by AI bots, and fast page response indicates a well-maintained, '
        'accessible site — a quality signal for AI platforms.'
    ))

    add_heading(doc, 'Page Load Time', level=2)

    if load_time is not None:
        perf_status = 'Fast (under 3 seconds — good)' if load_time_ok else 'Slow (over 3 seconds — needs improvement)'
        add_styled_para(doc, f'Measured load time: {load_time}s — {perf_status}')

        if not load_time_ok:
            add_callout_box(doc, 'Recommendation: Improve Page Load Time', [
                f'The page took {load_time}s to respond. AI crawlers have finite budgets '
                'and will deprioritise slow sites. Aim for under 3 seconds (ideally under 1s).',
                '',
                'Common fixes:',
                '• Enable server-side caching (Redis, Varnish, or CDN)',
                '• Compress images (WebP format)',
                '• Minify CSS and JavaScript',
                '• Use a Content Delivery Network (CDN) like Cloudflare',
                '• Enable GZIP/Brotli compression',
            ], 'C87814', 'FFF8E6')
    else:
        add_styled_para(doc, 'Load time could not be measured during this audit.')

    add_heading(doc, 'Core Web Vitals', level=2)

    add_styled_para(doc, (
        'This audit measures server response time only. For a complete performance picture, '
        'check Google Search Console (Core Web Vitals report) or PageSpeed Insights '
        '(pagespeed.web.dev) for LCP, INP, and CLS scores. Google AI Overview prioritises '
        'sites with good Core Web Vitals scores.'
    ))

    add_table(doc,
              ['Metric', 'What it measures', 'Good threshold'],
              [
                  ['LCP (Largest Contentful Paint)', 'Main content load speed', '< 2.5 seconds'],
                  ['INP (Interaction to Next Paint)', 'Interactivity responsiveness', '< 200ms'],
                  ['CLS (Cumulative Layout Shift)', 'Visual stability', '< 0.1'],
                  ['TTFB (Time to First Byte)', 'Server response speed', '< 800ms'],
              ])

    doc.add_page_break()

    # ── 5. DOMAIN AUTHORITY ──────────────────────────────────────────────────
    add_heading(doc, '5. Domain Authority & Backlinks', level=1)

    add_styled_para(doc, (
        'Domain authority is a significant factor in AI citation decisions. AI engines '
        'preferentially cite sources with high authority — indicated by the number of '
        'referring domains, total backlinks, and domain rank. Sites with strong backlink '
        'profiles are more likely to be included in AI-generated answers.'
    ))

    if backlinks_rank is not None or referring_domains is not None or total_backlinks is not None:
        add_heading(doc, 'Domain Metrics (from DataForSEO Backlinks API)', level=2)
        bl_rows = []
        if backlinks_rank is not None:
            bl_rows.append(['Domain Rank', str(backlinks_rank),
                            'Strong' if backlinks_rank and backlinks_rank >= 30 else 'Moderate' if backlinks_rank and backlinks_rank >= 10 else 'Weak'])
        if referring_domains is not None:
            bl_rows.append(['Referring Domains', f'{referring_domains:,}',
                            'Strong' if referring_domains >= 500 else 'Moderate' if referring_domains >= 50 else 'Weak'])
        if total_backlinks is not None:
            bl_rows.append(['Total Backlinks', f'{total_backlinks:,}',
                            'Strong' if total_backlinks >= 5000 else 'Moderate' if total_backlinks >= 200 else 'Weak'])
        add_table(doc, ['Metric', 'Value', 'Assessment'], bl_rows)

        if referring_domains is not None and referring_domains < 50:
            add_callout_box(doc, 'Recommendation: Build Domain Authority', [
                f'{client_domain} has {referring_domains:,} referring domains. '
                'Domains with fewer than 50 referring domains tend to receive lower AI citation rates. '
                '',
                'To improve domain authority:',
                '• Publish original research, statistics, or data that others will cite',
                '• Build relationships with industry publications for guest posts and mentions',
                '• Create linkable assets: free tools, templates, guides',
                '• Register in industry directories and association listings',
                '• Submit to relevant press releases and PR platforms',
            ], 'C87814', 'FFF8E6')
    else:
        add_styled_para(doc, (
            'Domain authority metrics were not available during this audit. '
            'Check Ahrefs, Moz, or SEMrush for Domain Rating / Domain Authority scores. '
            'For AI citation prioritisation, a Domain Rating of 40+ is typically a '
            'meaningful threshold.'
        ))

    add_heading(doc, 'Why Backlinks Matter for GEO', level=2)

    add_table(doc,
              ['Authority Level', 'Referring Domains', 'ChatGPT Cite Probability', 'Priority'],
              [
                  ['High Authority', '500+', 'High — frequently cited', 'Already strong'],
                  ['Medium Authority', '50–500', 'Moderate — cited for niche queries', 'Build further'],
                  ['Low Authority', '< 50', 'Low — rarely cited unprompted', 'Priority gap'],
              ])

    doc.add_page_break()

    # ── 6. CONTENT ANALYSIS ─────────────────────────────────────────────────
    add_heading(doc, '6. Content Analysis: Princeton GEO Methods', level=1)  # noqa

    add_styled_para(doc, (
        'The Princeton GEO framework identifies 9 content methods ranked by their impact on '
        'AI visibility. This section assesses how well the site\'s content applies these methods, '
        'based on what was accessible during the crawl.'
    ))

    add_styled_para(doc, (
        'Note: This automated audit assesses technical signals only. A full content audit '
        'requires manual review of key pages against the Princeton GEO methods. '
        'Refer to the Content Guide for detailed editorial guidance.'
    ), italic=True, color=GRAY)

    # Technical GEO signals we can actually detect
    add_heading(doc, 'Detected GEO Signals', level=2)

    geo_rows = []

    # Schema → cite sources / structured data
    geo_rows.append(['Cite Sources / Structured Data', '+40%',
                     'Partial' if jsonld_count > 0 else 'No',
                     f'{jsonld_count} JSON-LD block{"s" if jsonld_count != 1 else ""} detected'
                     if jsonld_count > 0 else 'No JSON-LD schema found — major gap'])

    # Meta signals → authoritative tone (proxy)
    title_signal = 'Moderate' if title else 'No'
    geo_rows.append(['Authoritative Tone', '+25%', title_signal,
                     'Title and meta present — check content for confident, expert language'
                     if title else 'Missing title undermines authority signals'])

    # Load time → fluency / accessibility proxy
    geo_rows.append(['Fluency / Accessibility', '+15–30%',
                     'Moderate' if load_time_ok else 'Weak',
                     f'Page loads in {load_time}s' if load_time else 'Could not measure'])

    # H1 → easy to understand
    geo_rows.append(['Easy-to-Understand Structure', '+20%',
                     'Moderate' if h1 else 'No',
                     'H1 heading present — check for clear, jargon-free writing'
                     if h1 else 'No H1 — basic content structure signals missing'])

    # What we cannot detect automatically
    geo_rows.append(['Statistics Addition', '+37%', 'Unknown',
                     'Manual review required — check pages for specific numbers and data points'])
    geo_rows.append(['Quotation Addition', '+30%', 'Unknown',
                     'Manual review required — check pages for expert quotes with attribution'])
    geo_rows.append(['Keyword Stuffing (avoid)', '−10%', 'Unknown',
                     'Manual review required — check for unnatural keyword repetition'])

    add_table(doc, ['GEO Method', 'Boost', 'Detected', 'Notes'], geo_rows)

    add_heading(doc, 'Content Recommendations', level=2)

    add_callout_box(doc, 'Priority: Add Statistics and Data Points (+37% boost)', [
        'Include specific, factual numbers throughout key pages. Examples:',
        '• Years in business, number of clients, project counts',
        '• Awards, accreditations, qualifications held',
        '• Customer satisfaction scores or testimonials with specifics',
        '• Pricing ranges, timelines, quantities',
        'Even approximate data is better than none. '
        '"Over 200 projects delivered" is far more citable than "extensive experience."'
    ], '167832', 'EBF9F0')

    add_callout_box(doc, 'Priority: Add Expert Quotes (+30% boost)', [
        'Include named quotes from team members, clients, or industry bodies. Format:',
        '"[Quote text]" — Name, Role/Organisation',
        '',
        'AI engines treat attributed quotes as high-quality, citable content. '
        'Anonymous testimonials have less impact. '
        'Client quotes are particularly valuable — they provide social proof and '
        'factual evidence that AI engines can extract and cite.'
    ], '167832', 'EBF9F0')

    add_callout_box(doc, 'Priority: Create FAQ Content (+40% boost)', [
        'Add a FAQ section to key pages answering common questions. Structure them as:',
        'Q: [Common question your audience would ask an AI]',
        'A: [Direct, factual, concise answer]',
        '',
        'FAQ content with FAQPage schema provides the single highest AI visibility boost '
        '(+40%) in the Princeton research. Target the exact questions people ask AI engines '
        'about your product, service, or area of expertise.'
    ], '167832', 'EBF9F0')

    doc.add_page_break()

    # ── 6. AI PLATFORM ASSESSMENT ───────────────────────────────────────────
    add_heading(doc, '7. AI Platform Assessment', level=1)

    add_styled_para(doc, (
        'Different AI platforms use different signals and indexes. Understanding each platform\'s '
        'requirements helps prioritise GEO improvements for maximum reach.'
    ))

    add_table(doc,
              ['Platform', 'Primary Index', 'Key Factor', f'{client_domain} Status'],
              [
                  ['ChatGPT', 'Bing / Web', 'Domain Authority + Content Quality',
                   'Good' if score >= 71 else ('Moderate' if score >= 41 else 'Weak')],
                  ['Perplexity', 'Own + Google', 'FAQ Schema + Semantic Relevance',
                   'Good' if jsonld_count > 0 and ai_bots else ('Moderate' if jsonld_count > 0 or ai_bots else 'Weak')],
                  ['Google AI Overview', 'Google', 'E-E-A-T + Structured Data',
                   'Good' if score >= 71 else ('Moderate' if score >= 41 else 'Weak')],
                  ['Microsoft Copilot', 'Bing', 'Bing Index + Entity Clarity',
                   'Moderate' if robots_exists else 'Weak'],
                  ['Claude', 'Brave Search', 'Factual Density + Source Authority',
                   'Good' if jsonld_count > 0 else ('Moderate' if title and description else 'Weak')],
              ])

    add_heading(doc, 'ChatGPT Citation Factors', level=2)

    add_table(doc,
              ['Factor', 'Weight', 'Status'],
              [
                  ['Domain Authority & Backlinks', '40%', 'Unknown — not measured in this audit'],
                  ['Content Quality & Utility', '35%',
                   'Moderate' if title and description and h1 else 'Weak'],
                  ['Platform Trust (Wikipedia, etc.)', '25%', 'Unknown'],
                  ['Content Recency', '3.2× boost', 'Unknown — check content freshness'],
                  ['Content-Answer Fit', 'Key factor',
                   'Moderate' if jsonld_count > 0 else 'Weak — needs FAQ-style content'],
              ])

    add_heading(doc, 'Google AI Overview Factors', level=2)

    add_table(doc,
              ['Factor', 'Status'],
              [
                  ['E-E-A-T (Experience, Expertise, Authority, Trust)',
                   'Moderate' if title and description else 'Weak'],
                  ['Structured data (JSON-LD)',
                   f'{jsonld_count} block{"s" if jsonld_count != 1 else ""} — {"good" if jsonld_count > 0 else "missing"}'],
                  ['Topical authority', 'Unknown — requires content depth review'],
                  ['Page speed', f'{load_time}s — {"fast" if load_time_ok else "slow"}' if load_time else 'Unknown'],
                  ['Knowledge Graph inclusion', 'Unknown — check Google Search Console'],
              ])

    doc.add_page_break()

    # ── 7. SCORECARD ────────────────────────────────────────────────────────
    add_heading(doc, '8. GEO Readiness Scorecard', level=1)

    scorecard_rows = []

    # Title
    t_score = 15 if title else 0
    scorecard_rows.append(['Page Title', f'{t_score}/15',
                            'Good' if title and audit.get('title_ok') else ('Weak' if title else 'Missing'),
                            'Present and correct length' if (title and audit.get('title_ok'))
                            else ('Present but too long' if title else 'Missing — critical gap')])

    # Description
    d_score = 10 if description else 0
    scorecard_rows.append(['Meta Description', f'{d_score}/10',
                            'Good' if description and audit.get('description_ok') else ('Weak' if description else 'Missing'),
                            'Present and good length' if (description and audit.get('description_ok'))
                            else ('Present but too long' if description else 'Missing')])

    # OG
    og_score = 5 if og_tags else 0
    scorecard_rows.append(['OG / Social Tags', f'{og_score}/5',
                            'Good' if og_tags else 'Missing',
                            'Present' if og_tags else 'Not detected'])

    # H1
    h1_score = 10 if h1 else 0
    scorecard_rows.append(['H1 Heading', f'{h1_score}/10',
                            'Good' if h1 else 'Missing',
                            'Present' if h1 else 'Missing — weakens page structure'])

    # Schema
    schema_score = 20 if jsonld_count > 0 else 0
    scorecard_rows.append(['JSON-LD Schema', f'{schema_score}/20',
                            'Good' if jsonld_count > 0 else 'Missing',
                            f'{jsonld_count} block{"s" if jsonld_count != 1 else ""}'
                            if jsonld_count > 0 else 'No structured data — major gap'])

    # AI bots
    bots_score = 15 if ai_bots else 0
    if ai_bots and ai_bots_blocked:
        bots_detail = f'{len(ai_bots)} allowed, {len(ai_bots_blocked)} blocked'
        bots_rating = 'Partial'
    elif ai_bots:
        bots_detail = f'{len(ai_bots)} bot{"s" if len(ai_bots) != 1 else ""} explicitly allowed'
        bots_rating = 'Good'
    elif ai_bots_blocked:
        bots_detail = f'{len(ai_bots_blocked)} bot{"s" if len(ai_bots_blocked) != 1 else ""} explicitly blocked'
        bots_rating = 'Blocked'
        bots_score = 0
    else:
        bots_detail = 'No AI bot directives in robots.txt'
        bots_rating = 'Missing'
    scorecard_rows.append(['AI Bot Access', f'{bots_score}/15', bots_rating, bots_detail])

    # Sitemap
    sitemap_score = 10 if has_sitemap else 0
    scorecard_rows.append(['XML Sitemap', f'{sitemap_score}/10',
                            'Good' if has_sitemap else 'Missing',
                            'Found at /sitemap.xml' if has_sitemap else 'Not found'])

    # Load time
    lt_score = 15 if load_time_ok else 0
    scorecard_rows.append(['Page Load Time', f'{lt_score}/15',
                            'Good' if load_time_ok else 'Weak',
                            f'{load_time}s' if load_time else 'Not measured'])

    add_table(doc, ['Category', 'Score', 'Rating', 'Detail'], scorecard_rows)

    # Overall score
    p = doc.add_paragraph()
    run = p.add_run(f'Overall GEO Score: {score} / 100')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = SCORE_COLOR
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)

    # Score legend
    legend_rows = [
        ['71–100', 'GEO Ready', 'Strong technical foundation. Focus on content enrichment.'],
        ['41–70', 'Needs Work', 'Some foundations in place. Address priority gaps systematically.'],
        ['0–40', 'Poor', 'Significant technical gaps. Prioritise schema, meta tags, and bot access.'],
    ]
    add_table(doc, ['Score', 'Rating', 'Meaning'], legend_rows)

    doc.add_page_break()

    # ── 8. RECOMMENDATIONS ──────────────────────────────────────────────────
    add_heading(doc, '9. Recommendations', level=1)

    rec_num = 1

    # Priority 0 — Critical blockers only if needed
    critical = []
    if not title:
        critical.append(('Add a page title', 'Every page must have a unique <title> tag. '
                         'This is the most basic SEO/GEO requirement and is currently missing.'))
    if not h1:
        critical.append(('Add an H1 heading', 'Every page needs exactly one H1 heading that '
                         'clearly describes the page topic.'))
    if jsonld_count == 0:
        critical.append(('Implement JSON-LD schema markup', 'No structured data was detected. '
                         'Add at minimum an Organization/LocalBusiness schema to every page, '
                         'and FAQPage schema to key content pages.'))

    if critical:
        add_heading(doc, 'Priority 0 — Critical (Must Fix First)', level=2)
        for title_text, body in critical:
            add_callout_box(doc, f'Recommendation {rec_num}: {title_text}', [body],
                            'B42828', 'FFF0F0')
            rec_num += 1

    # Priority 1 — High impact
    add_heading(doc, 'Priority 1 — High Impact', level=2)

    if not description:
        add_callout_box(doc, f'Recommendation {rec_num}: Write meta descriptions for all pages', [
            'Add a 120–155 character meta description to every page. Include the primary keyword '
            'and a clear value proposition. AI engines use meta descriptions as summary signals.'
        ], '0072BB', 'EDF6FC')
        rec_num += 1

    if not og_tags:
        add_callout_box(doc, f'Recommendation {rec_num}: Add Open Graph tags', [
            'Implement og:title, og:description, og:url, og:image on all pages. '
            'These are used by social platforms and AI engines for rich content interpretation.'
        ], '0072BB', 'EDF6FC')
        rec_num += 1

    if not ai_bots:
        add_callout_box(doc, f'Recommendation {rec_num}: Configure AI bot access in robots.txt', [
            'Add explicit Allow directives for GPTBot, ClaudeBot, PerplexityBot, and '
            'ChatGPT-User. This signals that AI crawlers are welcome and may improve '
            'crawl priority for your content.'
        ], '0072BB', 'EDF6FC')
        rec_num += 1

    if not has_sitemap:
        add_callout_box(doc, f'Recommendation {rec_num}: Create an XML sitemap', [
            'Generate a sitemap.xml and reference it in robots.txt. '
            'This helps AI crawlers discover all your pages efficiently.'
        ], '0072BB', 'EDF6FC')
        rec_num += 1

    add_callout_box(doc, f'Recommendation {rec_num}: Add FAQPage schema to key pages', [
        'Implement FAQPage JSON-LD on your most important content pages. '
        'Structure 4–8 Q&A pairs per page answering the questions your audience '
        'asks AI engines about your product or service. '
        'This provides the single highest AI visibility boost (+40%) in the Princeton research.'
    ], '0072BB', 'EDF6FC')
    rec_num += 1

    # Priority 2 — Content
    add_heading(doc, 'Priority 2 — Content Enrichment', level=2)

    add_callout_box(doc, f'Recommendation {rec_num}: Add statistics to key pages (+37% boost)', [
        'Enrich content with specific, verifiable numbers. '
        'Years in business, client counts, project counts, awards, accreditations. '
        'Statistics are the second highest-impact GEO method in the Princeton research.'
    ], '167832', 'EBF9F0')
    rec_num += 1

    add_callout_box(doc, f'Recommendation {rec_num}: Add expert and client quotes (+30% boost)', [
        'Include named, attributed quotes on product/service pages and the homepage. '
        'Format: "[Quote]" — Name, Title/Organisation. '
        'Client testimonials with specifics ("X saved us Y hours per week") are particularly citable.'
    ], '167832', 'EBF9F0')
    rec_num += 1

    add_callout_box(doc, f'Recommendation {rec_num}: Structure content for answer extraction', [
        'Rewrite key pages in an "answer-first" format: lead with the key fact or definition, '
        'then expand with supporting detail. AI engines prefer content structured as direct '
        'answers to questions over marketing copy or narrative prose.'
    ], '167832', 'EBF9F0')
    rec_num += 1

    # Priority 3 — Technical
    add_heading(doc, 'Priority 3 — Technical Improvements', level=2)

    if not load_time_ok and load_time:
        add_callout_box(doc, f'Recommendation {rec_num}: Improve page load time', [
            f'The page took {load_time}s to load. AI crawlers deprioritise slow sites. '
            'Implement server-side caching, image compression (WebP), and a CDN. '
            'Aim for under 1 second TTFB.'
        ], 'C87814', 'FFF8E6')
        rec_num += 1

    add_callout_box(doc, f'Recommendation {rec_num}: Implement IndexNow for Bing/Copilot', [
        'IndexNow notifies Bing instantly when new content is published, '
        'ensuring Microsoft Copilot indexes fresh content quickly. '
        'Most CMS platforms have IndexNow plugins available.'
    ], 'C87814', 'FFF8E6')
    rec_num += 1

    add_callout_box(doc, f'Recommendation {rec_num}: Register in Google Search Console', [
        'Verify the site in Google Search Console to monitor Core Web Vitals, '
        'crawl errors, and schema validation. Check the Rich Results Test for '
        'structured data issues: search.google.com/test/rich-results'
    ], 'C87814', 'FFF8E6')

    doc.add_page_break()

    # ── APPENDIX A ──────────────────────────────────────────────────────────
    add_heading(doc, 'Appendix A: Princeton GEO Methods Reference', level=1)

    add_table(doc,
              ['Method', 'AI Visibility Boost', 'Description'],
              [
                  ['Cite Sources', '+40%', 'Add authoritative citations and external references'],
                  ['Statistics Addition', '+37%', 'Include specific numbers and data points'],
                  ['Quotation Addition', '+30%', 'Add expert quotes with clear attribution'],
                  ['Authoritative Tone', '+25%', 'Use confident, definitive, expert language'],
                  ['Easy-to-understand', '+20%', 'Simplify complex concepts — plain English first'],
                  ['Technical Terms', '+18%', 'Include relevant domain-specific terminology'],
                  ['Unique Words', '+15%', 'Increase vocabulary diversity — avoid repetition'],
                  ['Fluency Optimization', '+15–30%', 'Improve readability and logical flow'],
                  ['Keyword Stuffing', '−10%', 'AVOID — unnatural repetition hurts AI visibility'],
              ])

    add_styled_para(doc, (
        'Source: Princeton University, IIT Delhi, Georgia Tech, Allen Institute for AI. '
        '"GEO: Generative Engine Optimization" (arXiv:2311.09735, KDD 2024).'
    ), italic=True, color=GRAY, size=9)

    # ── APPENDIX B ──────────────────────────────────────────────────────────
    add_heading(doc, 'Appendix B: AI Platform Quick Reference', level=1)

    add_table(doc,
              ['Platform', 'Primary Index', 'Key Requirement', 'Unique Factor'],
              [
                  ['ChatGPT', 'Bing / Web', 'Domain Authority', 'Content-Answer Fit'],
                  ['Perplexity', 'Own + Google', 'FAQ Schema', 'PDF indexing, semantic search'],
                  ['Google AI Overview', 'Google', 'E-E-A-T signals', 'Knowledge Graph + schema'],
                  ['Microsoft Copilot', 'Bing', 'Bing Webmaster', 'IndexNow for freshness'],
                  ['Claude', 'Brave Search', 'Factual density', 'Brave Search indexing'],
              ])

    # ── APPENDIX C ──────────────────────────────────────────────────────────
    add_heading(doc, 'Appendix C: Audit Data', level=1)

    add_styled_para(doc, 'Data collected during this automated audit:')

    add_table(doc,
              ['Data Point', 'Value', 'Method'],
              [
                  ['URL audited', url, 'Direct crawl'],
                  ['Page title', title or 'Not found', 'HTML <title> tag extraction'],
                  ['Title length', f'{audit.get("title_length", 0)} characters', 'Character count'],
                  ['Meta description', (description[:80] + '…' if len(description) > 80 else description) or 'Not found',
                   'HTML meta[name=description] extraction'],
                  ['H1 heading', h1 or 'Not found', 'HTML <h1> extraction'],
                  ['OG tags', _yes_no(og_tags), 'HTML meta[property=og:*] detection'],
                  ['JSON-LD blocks', str(jsonld_count), 'application/ld+json script detection'],
                  ['Load time', f'{load_time}s' if load_time else 'Not measured', 'urllib HTTP request timer'],
                  ['robots.txt', 'Found' if robots_exists else 'Not found', 'GET /robots.txt'],
                  ['AI bots allowed', ', '.join(ai_bots) or 'None', 'robots.txt Allow directive parsing'],
                  ['AI bots blocked', ', '.join(ai_bots_blocked) or 'None', 'robots.txt Disallow directive parsing'],
                  ['sitemap.xml', f'Found at {sitemap_url}' if has_sitemap else 'Not found', 'Multi-location sitemap detection'],
                  ['Domain Rank', str(backlinks_rank) if backlinks_rank is not None else 'Not fetched', 'DataForSEO Backlinks API'],
                  ['Referring Domains', f'{referring_domains:,}' if referring_domains is not None else 'Not fetched', 'DataForSEO Backlinks API'],
                  ['Total Backlinks', f'{total_backlinks:,}' if total_backlinks is not None else 'Not fetched', 'DataForSEO Backlinks API'],
                  ['GEO Score', f'{score}/100', 'Weighted scoring algorithm'],
                  ['Audit date', date_str, 'Report generation date'],
              ])

    return doc
