"""
Shared python-docx utility functions extracted from the original generators.
Import these in all report generator modules instead of duplicating them.
"""
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Brand colours
BLACK = RGBColor(26, 26, 26)
DARK = RGBColor(40, 40, 40)
GRAY = RGBColor(100, 100, 100)
LIGHT_GRAY = RGBColor(160, 160, 160)
WHITE = RGBColor(255, 255, 255)
RED = RGBColor(180, 40, 40)
GREEN = RGBColor(22, 120, 50)
AMBER = RGBColor(180, 120, 20)
RULE_BLUE = RGBColor(0, 114, 187)
TIP_AMBER = RGBColor(200, 150, 30)


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, side, color_hex, size='12', style='single'):
    """Set a specific border on a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tc_pr.append(borders)
    border_el = parse_xml(
        f'<w:{side} {nsdecls("w")} w:val="{style}" w:sz="{size}" '
        f'w:space="0" w:color="{color_hex}"/>'
    )
    existing = borders.find(qn(f'w:{side}'))
    if existing is not None:
        borders.remove(existing)
    borders.append(border_el)


def remove_cell_borders(cell):
    """Remove all borders from a cell."""
    for side in ('top', 'bottom', 'left', 'right'):
        set_cell_border(cell, side, 'FFFFFF', size='0', style='none')


def add_styled_para(doc, text, size=10, bold=False, italic=False, color=None,
                    alignment=None, space_after=None, space_before=None):
    """Add a paragraph with specific styling."""
    if color is None:
        color = DARK
    if space_after is None:
        space_after = Pt(6)
    if space_before is None:
        space_before = Pt(0)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    return p


def add_heading(doc, text, level=1):
    """Add a heading with Numiko styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLACK
        run.font.name = 'Calibri'
    return h


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
        run = p.add_run(text)
    else:
        run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        set_cell_shading(cell, '1A1A1A')

    for row_idx, row in enumerate(rows):
        for col_idx, val in enumerate(row):
            if col_idx >= len(headers):
                break
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK
            if row_idx % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')

    doc.add_paragraph()
    return table


def add_callout_box(doc, heading, body_lines, accent_hex, bg_hex, heading_color=None):
    """Add a callout box as a single-cell table with left border accent."""
    if heading_color is None:
        heading_color = RGBColor(
            int(accent_hex[0:2], 16),
            int(accent_hex[2:4], 16),
            int(accent_hex[4:6], 16),
        )

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_hex)
    set_cell_border(cell, 'left', accent_hex, size='24')
    set_cell_border(cell, 'top', bg_hex, size='4')
    set_cell_border(cell, 'bottom', bg_hex, size='4')
    set_cell_border(cell, 'right', bg_hex, size='4')

    p = cell.paragraphs[0]
    run = p.add_run(heading)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = heading_color
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

    for line in body_lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()
    return table


def add_tip_box(doc, text, heading='Tip'):
    """Amber-accented tip box."""
    return add_callout_box(doc, heading, [text], 'C89614', 'FFF8E6',
                           heading_color=TIP_AMBER)


def add_example_box(doc, before_text, after_text, annotations=None):
    """Before/after comparison box."""
    add_callout_box(doc, 'BEFORE \u2014 Weak content', [before_text],
                    'B42828', 'FFF0F0', heading_color=RED)
    lines = [after_text]
    if annotations:
        lines.append('')
        lines.append('GEO methods applied: ' + annotations)
    add_callout_box(doc, 'AFTER \u2014 GEO-optimised content', lines,
                    '167832', 'EBFaF0', heading_color=GREEN)


def add_golden_rule(doc, number, title, boost, description):
    """Numbered golden rule box."""
    heading = f'{number}. {title}    {boost}'
    add_callout_box(doc, heading, [description], '0072BB', 'EDF6FC',
                    heading_color=RULE_BLUE)


def add_checklist_item(doc, text):
    """Checkbox-style item."""
    p = doc.add_paragraph()
    run = p.add_run('[ ]  ')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    return p
