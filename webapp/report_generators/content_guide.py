"""
Parameterized Content Guide generator.
Produces a GEO/SEO content guide DOCX customised per client.

params dict keys:
  client_name   - e.g. "GCHQ"              (used on cover + examples)
  client_domain - e.g. "gchq.gov.uk"       (used on cover)
  project_name  - e.g. "Stillwater Project" (optional, used on cover)
  date          - e.g. "February 2026"      (used on cover)
  cms           - e.g. "Drupal"             (used in section 3 headings)
  logo_path     - absolute path to logo image (optional, falls back to agency logo)
"""
import os
import logging
from datetime import datetime
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx import Document

from .docx_helpers import (
    BLACK, DARK, GRAY, LIGHT_GRAY, WHITE, RED, GREEN, RULE_BLUE, TIP_AMBER,
    add_styled_para, add_heading, add_bullet, add_table,
    add_callout_box, add_tip_box, add_example_box, add_golden_rule, add_checklist_item,
    create_document,
)

logger = logging.getLogger(__name__)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DEFAULT_LOGO = os.path.join(BASE_DIR, 'fonts', 'numiko_logo.png')


def build_content_guide(params: dict) -> Document:
    client_name = params.get('client_name') or 'Client'
    client_domain = params.get('client_domain') or ''
    project_name = params.get('project_name') or ''
    date_str = params.get('date') or datetime.now().strftime('%B %Y')
    cms = params.get('cms') or 'Drupal'
    logo_path = params.get('logo_path') or DEFAULT_LOGO

    doc = create_document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, height=Cm(3))

    doc.add_paragraph()
    add_styled_para(doc, 'Content Guide', size=28, bold=True, color=BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_styled_para(doc, 'SEO & GEO', size=22, bold=True, color=BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_styled_para(doc, 'Writing for Search Engines and AI', size=14, color=GRAY,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))

    cover_subtitle = client_domain
    if project_name:
        cover_subtitle = f'{client_domain} \u2014 {project_name}' if client_domain else project_name
    if cover_subtitle:
        add_styled_para(doc, cover_subtitle, size=11, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    add_styled_para(doc, date_str, size=10, color=LIGHT_GRAY,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))
    add_styled_para(doc,
                    f'A practical guide for content authors{" at " + client_name if client_name else ""}',
                    size=8, color=LIGHT_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=Pt(2))
    add_styled_para(doc, 'Based on Princeton GEO research (KDD 2024)',
                    size=8, color=LIGHT_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ── 1. WHY THIS MATTERS ──────────────────────────────────────────────────
    add_heading(doc, '1. Why This Matters', level=1)

    add_styled_para(doc, (
        'People are changing how they find information. Instead of typing a search into Google '
        'and clicking through results, millions now ask AI assistants directly: ChatGPT, '
        'Perplexity, Google AI Overview, Microsoft Copilot, and Claude. These AI engines '
        'read website content and quote the best sources in their answers.'
    ))

    add_styled_para(doc, 'Your goal as a content author: make your content the source that AI engines cite.',
                    bold=True, color=BLACK)

    add_styled_para(doc, (
        'This is called Generative Engine Optimization (GEO). Think of it as the next step '
        'beyond SEO. With traditional SEO, you wanted to rank on the first page of Google. '
        'With GEO, you want to be the answer that AI gives when someone asks a question.'
    ))

    add_heading(doc, 'The Opportunity', level=2)
    if client_name and client_domain:
        add_styled_para(doc, (
            f'Making {client_domain} visible to AI crawlers is only the first step. Visibility '
            f'alone is not enough. AI engines only cite content that is well-structured, '
            f'evidence-based, and directly answers questions. That is where you come in.'
        ))
    else:
        add_styled_para(doc, (
            'Making your site visible to AI crawlers is only the first step. Visibility '
            'alone is not enough. AI engines only cite content that is well-structured, '
            'evidence-based, and directly answers questions. That is where you come in.'
        ))

    add_heading(doc, 'What AI Engines Look For', level=2)
    add_styled_para(doc, (
        'Princeton University research (published 2024) studied what makes AI engines cite '
        'one page over another. They found 9 specific methods that increase visibility by '
        'up to 40%. We have distilled these into 5 practical rules you can apply to every '
        'page you write.'
    ))

    add_table(doc,
              ['What AI Engines Want', 'What They Avoid'],
              [
                  ['Direct answers to questions', 'Vague introductions'],
                  ['Statistics and data points', 'Claims without evidence'],
                  ['Expert quotes with attribution', 'Anonymous assertions'],
                  ['Links to authoritative sources', 'Unsupported statements'],
                  ['Confident, expert language', 'Marketing fluff'],
                  ['Clear definitions', 'Jargon without explanation'],
              ])

    add_tip_box(doc, (
        'Every page you write is competing with Wikipedia, government reports, and news '
        'articles to be the source AI cites. The more evidence and structure you provide, '
        'the more likely your content will be chosen.'
    ))

    doc.add_page_break()

    # ── 2. THE 5 GOLDEN RULES ────────────────────────────────────────────────
    add_heading(doc, '2. The 5 Golden Rules', level=1)
    add_styled_para(doc, (
        'These five rules are based on the Princeton GEO research. Apply them to every page '
        'you write. Each rule shows the percentage boost it gives to your chance of being '
        'cited by AI engines.'
    ))

    add_golden_rule(doc, 1, 'Answer the question first', '+20% to +30%', (
        'Lead with a clear definition or direct answer in the first two sentences. Write as '
        'if someone has just asked \u201cWhat is [your topic]?\u201d The opening paragraph should give '
        'them the answer immediately \u2014 not a welcome message, not background context, not a '
        'mission statement. The definition comes first, the detail follows. AI engines extract '
        'the first few sentences as their answer. Make those sentences count.'
    ))

    add_golden_rule(doc, 2, 'Back it up with numbers', '+37%', (
        'Include at least two statistics or data points on every page. Numbers make content '
        '37% more likely to be cited by AI. Examples: staff numbers, incident counts, dates, '
        'percentages, budgets, programme participation figures, timelines. Even approximate or '
        'publicly-available figures are vastly better than none.'
    ))

    add_golden_rule(doc, 3, 'Quote the experts', '+30%', (
        'Add at least one attributed quote per page. Good sources: leadership statements, '
        'oversight reports, partner organisation leaders, employee testimonials, academic experts, '
        'government reviews. Format: \u201cQuote text\u201d \u2014 Name, Role, Organisation. Quotes add '
        'authority, break up text, and give AI engines a human voice to cite.'
    ))

    add_golden_rule(doc, 4, 'Cite your sources', '+40%', (
        'Link to authoritative external references \u2014 at least three per page. This is the '
        'single highest-impact method, boosting AI visibility by 40%. Good citation sources: '
        'government reports, legislation, partner organisations, academic research, and '
        'reputable news coverage. Use inline links in rich text. Every claim should have a source.'
    ))

    add_golden_rule(doc, 5, 'Write like an authority, not a brochure', '+25%', (
        'Use confident, definitive language. Include technical terms but define them on first '
        'use. Avoid marketing fluff, vague aspirations, and repetitive slogans. Vary your '
        'vocabulary: unique words boost visibility by 15%.'
    ))

    add_callout_box(doc, 'Anti-Rule: Never Stuff Keywords (\u221210%)', [
        'Repeating the same keyword or phrase hurts your AI visibility. AI engines actively '
        'penalise repetitive content. Instead, use varied phrasing that conveys the same meaning '
        'in different ways.'
    ], 'B42828', 'FFF0F0', heading_color=RED)

    doc.add_page_break()

    # ── 3. CMS CONTENT TOOLKIT ───────────────────────────────────────────────
    add_heading(doc, f'3. Your {cms} Content Toolkit', level=1)
    add_styled_para(doc, (
        f'The {cms} CMS gives you specific content components that are designed to help with '
        f'SEO and GEO. Here are the five components you will use most often, and when to '
        f'reach for each one.'
    ))

    add_heading(doc, 'FAQ / Accordion Component     +40%', level=3)
    add_styled_para(doc, 'Ask yourself: \u201cWould a reader ask a question about this topic?\u201d',
                    italic=True, color=GRAY)
    add_styled_para(doc, (
        'FAQ or accordion components create expandable question-and-answer sections. Behind the '
        'scenes, these are marked up as FAQ schema \u2014 a special code that tells AI engines '
        '\u201cthis is a direct answer to a question.\u201d Aim for at least 3 questions per page.'
    ))

    add_heading(doc, 'Quote Component     +30%', level=3)
    add_styled_para(doc, 'Ask yourself: \u201cDo I have a statement from a leader, expert, or partner?\u201d',
                    italic=True, color=GRAY)
    add_styled_para(doc, (
        'A quote component displays an attributed quotation with the speaker\u2019s name, role, and '
        'organisation. AI engines value expert quotes because they add human authority to '
        'factual content.'
    ))

    add_heading(doc, 'Statistics Component     +37%', level=3)
    add_styled_para(doc, 'Ask yourself: \u201cDo I have a number, metric, or data point?\u201d',
                    italic=True, color=GRAY)
    add_styled_para(doc, (
        'A statistics component presents data in a visually prominent format. Even approximate '
        'numbers are better than none. Specific figures are far more citable than vague claims.'
    ))

    add_heading(doc, 'Key Fact Component     +37%', level=3)
    add_styled_para(doc, 'Ask yourself: \u201cIs there a headline data point that deserves prominence?\u201d',
                    italic=True, color=GRAY)
    add_styled_para(doc, (
        'Similar to the statistics component but designed for single, standout facts. These '
        'give AI engines the definitive facts they need to answer questions accurately.'
    ))

    add_heading(doc, 'Rich Text / Body Content     +40%', level=3)
    add_styled_para(doc, 'Ask yourself: \u201cMain content area \u2014 where I write narrative with inline links\u201d',
                    italic=True, color=GRAY)
    add_styled_para(doc, (
        'Your primary writing area. The key GEO benefit here is inline links to authoritative '
        'sources. Every time you reference a report, a law, or a partner organisation, link to it. '
        'This tells AI engines that your content is evidence-based and trustworthy.'
    ))

    add_tip_box(doc, (
        'Don\u2019t skip the metadata fields. Every page has a Title and Description field. '
        'The Title appears in search results and AI citations (keep it under 60 characters). '
        'The Description is your pitch to both humans and AI (under 155 characters).'
    ), heading='Important: Metadata Fields')

    doc.add_page_break()

    # ── 4. PAGE-BY-PAGE PLAYBOOK ─────────────────────────────────────────────
    add_heading(doc, '4. Page-by-Page Playbook', level=1)
    add_styled_para(doc, (
        'This section gives specific guidance for different page types. For each type, you '
        'will find: what the page needs to achieve, a content checklist, and recommended components.'
    ))

    add_heading(doc, 'Home / About Pages', level=2)
    add_styled_para(doc, 'Goal: Define who you are and what you do, with evidence.', bold=True, color=BLACK)
    add_styled_para(doc, 'Opening sentence structure:', space_after=Pt(2))
    add_styled_para(doc,
                    f'"{client_name} is [definition]. Founded in [year], it [core mission]."'
                    if client_name else '"[Organisation] is [definition]. Founded in [year], it [core mission]."',
                    italic=True, color=GRAY)

    add_heading(doc, 'Content Checklist', level=3)
    add_bullet(doc, 'Opening definition of the organisation within the first 2 sentences')
    add_bullet(doc, 'Founding date and brief history')
    add_bullet(doc, 'Staff count and key locations')
    add_bullet(doc, 'Core missions or services (3\u20135 bullet points)')
    add_bullet(doc, 'Oversight or governance structure')
    add_bullet(doc, 'At least 1 quote from a leader or trusted source')
    add_bullet(doc, 'Links to: key reports, legislation, partner organisations')

    add_heading(doc, 'Recommended Components', level=3)
    add_bullet(doc, 'FAQ/Accordion: \u201cWhat is [organisation]?\u201d, \u201cWhat does it do?\u201d, \u201cWho oversees it?\u201d')
    add_bullet(doc, 'Quote: Leadership statement on mission')
    add_bullet(doc, 'Key Fact: Founded date, staff count, key locations')
    add_bullet(doc, 'Statistics: Key metrics and impact data')

    add_heading(doc, 'News / Publication Pages', level=2)
    add_styled_para(doc, 'Goal: Lead with the key finding, supported by data and expert commentary.',
                    bold=True, color=BLACK)
    add_styled_para(doc, 'Opening sentence structure:', space_after=Pt(2))
    add_styled_para(doc, (
        '"[Key finding or announcement] according to [source/date]. [Supporting statistic]."'
    ), italic=True, color=GRAY)

    add_heading(doc, 'Content Checklist', level=3)
    add_bullet(doc, 'Headline finding or announcement in the first sentence')
    add_bullet(doc, 'Date of publication or event')
    add_bullet(doc, 'Key statistic from the report or announcement')
    add_bullet(doc, 'Expert quote providing commentary')
    add_bullet(doc, 'Link to full publication or source document')
    add_bullet(doc, 'Context: why this matters, what it means for the audience')

    add_heading(doc, 'Career / Role Pages', level=2)
    add_styled_para(doc, 'Goal: Provide structured, specific role information that AI assistants can cite.',
                    bold=True, color=BLACK)
    add_styled_para(doc, 'Opening sentence structure:', space_after=Pt(2))
    add_styled_para(doc, (
        '"[Role Title] at [organisation] [1-sentence description]. '
        'Based in [locations], the role offers [key benefits]."'
    ), italic=True, color=GRAY)

    add_heading(doc, 'Content Checklist', level=3)
    add_bullet(doc, 'Role title and 1-sentence description in the opening')
    add_bullet(doc, 'Salary range (where publishable)')
    add_bullet(doc, 'Locations and working arrangements')
    add_bullet(doc, 'Key responsibilities (3\u20135 bullet points)')
    add_bullet(doc, 'Required qualifications and skills')
    add_bullet(doc, 'Benefits and development opportunities')
    add_bullet(doc, 'At least 1 employee testimonial or quote')

    doc.add_page_break()

    # ── 5. SEO & GEO CHECKLIST ───────────────────────────────────────────────
    add_heading(doc, '5. The SEO & GEO Checklist', level=1)
    add_styled_para(doc, (
        'Use this checklist for every page you write or rewrite. Tick off each item '
        'before you publish.'
    ))

    add_heading(doc, 'Structure & Metadata', level=2)
    add_checklist_item(doc, 'Page title: Clear, descriptive, under 60 characters')
    add_checklist_item(doc, 'Meta description: Compelling summary, under 155 characters')
    add_checklist_item(doc, 'H1 heading: One per page, matches what people search for')
    add_checklist_item(doc, 'URL: Clean and descriptive')

    add_heading(doc, 'Opening Content', level=2)
    add_checklist_item(doc, 'First sentence directly answers \u201cWhat is [topic]?\u201d')
    add_checklist_item(doc, 'Definition or key fact within the first 50 words')
    add_checklist_item(doc, 'No \u201cwelcome to\u201d or \u201cthis page is about\u201d preambles')

    add_heading(doc, 'GEO Requirements', level=2)
    add_checklist_item(doc, 'At least 2 statistics or data points')
    add_checklist_item(doc, 'At least 1 expert quote with attribution')
    add_checklist_item(doc, 'At least 3 links to authoritative external sources')
    add_checklist_item(doc, 'FAQ section with 3+ questions')
    add_checklist_item(doc, 'Technical terms are defined on first use')

    add_heading(doc, 'Content Quality', level=2)
    add_checklist_item(doc, 'No keyword or phrase repeated more than 3 times')
    add_checklist_item(doc, 'Content reads naturally when spoken aloud (fluency test)')
    add_checklist_item(doc, 'No marketing fluff or vague claims without evidence')
    add_checklist_item(doc, 'Vocabulary is varied (no repetitive phrasing)')
    add_checklist_item(doc, 'Content is current and up-to-date (no outdated references)')
    add_checklist_item(doc, 'All claims are supported by a citation or data point')

    add_tip_box(doc, (
        'The fluency test: read your content aloud. If it sounds like a brochure or a '
        'mission statement, rewrite it. If it sounds like an expert answering a question, '
        'you are on the right track.'
    ), heading='The Read-Aloud Test')

    doc.add_page_break()

    # ── 6. BEFORE & AFTER EXAMPLES ───────────────────────────────────────────
    add_heading(doc, '6. Before & After Examples', level=1)
    add_styled_para(doc, (
        'These examples show the difference between content that AI engines ignore and content '
        'that AI engines cite. Study the annotations to understand which GEO methods were applied.'
    ))

    add_heading(doc, 'Example 1: About Page Introduction', level=2)
    add_example_box(doc,
                    f'{client_name} works to help people. We are committed to making a difference. '
                    f'Our work is important and our people are dedicated.',
                    f'{client_name} is [definition in one sentence]. Founded in [year], '
                    f'{client_name} [core mission with specific scope]. With [staff count] staff '
                    f'across [locations], the organisation [key activities]. '
                    f'[Organisation] operates under [legislation/framework] '
                    f'and is accountable to [oversight body].',
                    'Definition first (+20%), founding date and staff count (+37%), '
                    'citations to legislation (+40%), technical terms defined (+18%), '
                    'authoritative tone (+25%)')

    add_heading(doc, 'Example 2: Statistics Usage', level=2)
    add_example_box(doc,
                    f'{client_name} handles many cases every year. We deal with a large volume '
                    f'of work and our team works hard to achieve good results.',
                    f'{client_name} handled [X] cases in [year], a [Y%] increase from the '
                    f'previous year (Annual Report [year]). Of these, [Z%] were resolved within '
                    f'[timeframe], exceeding the [benchmark] target set by [oversight body]. '
                    f'\u201c[Expert quote about results]\u201d \u2014 [Name, Role, Organisation].',
                    'Specific numbers (+37%), citation to annual report (+40%), '
                    'expert quote (+30%), benchmark context (+25%)')

    doc.add_page_break()

    # ── 7. QUICK REFERENCE CARD ──────────────────────────────────────────────
    add_heading(doc, '7. Quick Reference Card', level=1)
    add_styled_para(doc, 'Print this page and keep it beside your screen while writing.')

    add_heading(doc, 'The 5 Golden Rules', level=2)
    add_table(doc,
              ['Rule', 'What to Do', 'Boost'],
              [
                  ['1. Answer first', 'Definition or direct answer in the first 2 sentences', '+20\u201330%'],
                  ['2. Use numbers', 'At least 2 statistics or data points per page', '+37%'],
                  ['3. Quote experts', 'At least 1 attributed quote per page', '+30%'],
                  ['4. Cite sources', 'At least 3 links to authoritative references', '+40%'],
                  ['5. Be the authority', 'Confident language, define technical terms', '+25%'],
              ])

    add_heading(doc, 'Content Components', level=2)
    add_table(doc,
              ['Component', 'Use When...', 'GEO Boost'],
              [
                  ['FAQ/Accordion', 'A reader might ask a question about this', '+40%'],
                  ['Quote', 'You have an expert statement', '+30%'],
                  ['Statistics', 'You have numbers or metrics', '+37%'],
                  ['Key Fact', 'There is a headline data point', '+37%'],
                  ['Rich Text', 'Writing narrative with inline source links', '+40%'],
              ])

    add_heading(doc, 'Per-Page Minimums', level=2)
    add_table(doc,
              ['Requirement', 'Minimum', 'Component'],
              [
                  ['Statistics / data points', '2 per page', 'Statistics or Key Fact'],
                  ['Expert quotes', '1 per page', 'Quote component'],
                  ['External citations', '3 per page', 'Inline links in body text'],
                  ['FAQ questions', '3 per page', 'FAQ/Accordion component'],
                  ['Definition in opening', 'First 2 sentences', 'Opening body text'],
              ])

    add_callout_box(doc, 'The One Thing to Avoid', [
        'Keyword stuffing: repeating the same phrase more than 2\u20133 times on a page '
        'actively hurts your AI visibility by 10%. Vary your language.'
    ], 'B42828', 'FFF0F0', heading_color=RED)

    add_tip_box(doc, (
        'Remember: AI engines do not cite brochures. They cite sources of evidence. '
        'Every sentence should either state a fact, provide a statistic, quote an expert, '
        'or cite a source. If a sentence does none of these, ask yourself: does this '
        'sentence add value, or is it filler?'
    ), heading='The Bottom Line')

    return doc
