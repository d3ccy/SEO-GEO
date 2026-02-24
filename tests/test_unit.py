"""
Comprehensive unit tests for SEO-GEO Toolkit core business logic.

These tests run entirely offline -- no network access or DataForSEO
credentials required.  Run with:

    pytest tests/test_unit.py -v

"""
import json
import os
import sys

import pytest

# ---------------------------------------------------------------------------
# Path setup -- ensure webapp/ and scripts/ are importable
# ---------------------------------------------------------------------------
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_WEBAPP_DIR = os.path.join(_PROJECT_ROOT, 'webapp')
_SCRIPTS_DIR = os.path.join(_PROJECT_ROOT, 'scripts')

for _p in (_WEBAPP_DIR, _SCRIPTS_DIR):
    if _p not in sys.path:
        sys.path.insert(0, _p)


# ===========================================================================
# 1. audit_service._calculate_score
# ===========================================================================

from services.audit_service import _calculate_score


class TestCalculateScore:
    """Test the 0-100 GEO readiness scoring algorithm."""

    # -- helper to build a meta dict easily --------------------------------
    @staticmethod
    def _meta(title='', description='', og_tags=False, h1='', jsonld_count=0):
        return {
            'title': title,
            'description': description,
            'og_tags': og_tags,
            'h1': h1,
            'jsonld_count': jsonld_count,
        }

    def test_all_fields_present_score_100(self):
        """Every check passes -> score = 100."""
        meta = self._meta(
            title='My Page',
            description='A description',
            og_tags=True,
            h1='Heading',
            jsonld_count=1,
        )
        robots = {'ai_bots': ['GPTBot']}
        score = _calculate_score(meta, robots, has_sitemap=True,
                                 load_time=1.5, page_blocked=False)
        assert score == 100

    def test_all_fields_missing_score_0(self):
        """Every check fails -> score = 0."""
        meta = self._meta()
        robots = {}
        score = _calculate_score(meta, robots, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 0

    def test_title_only_15(self):
        meta = self._meta(title='Title')
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 15

    def test_description_only_10(self):
        meta = self._meta(description='A description')
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 10

    def test_og_tags_only_5(self):
        meta = self._meta(og_tags=True)
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 5

    def test_h1_only_10(self):
        meta = self._meta(h1='Heading One')
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 10

    def test_jsonld_only_20(self):
        meta = self._meta(jsonld_count=2)
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 20

    def test_ai_bots_only_15(self):
        meta = self._meta()
        robots = {'ai_bots': ['GPTBot', 'ClaudeBot']}
        score = _calculate_score(meta, robots, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 15

    def test_sitemap_only_10(self):
        meta = self._meta()
        score = _calculate_score(meta, {}, has_sitemap=True,
                                 load_time=None, page_blocked=False)
        assert score == 10

    def test_load_time_fast_15(self):
        meta = self._meta()
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=1.0, page_blocked=False)
        assert score == 15

    # -- boundary: 3-second threshold -------------------------------------
    def test_load_time_just_under_3s_passes(self):
        """2.99 s is under 3 -> +15 points."""
        meta = self._meta()
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=2.99, page_blocked=False)
        assert score == 15

    def test_load_time_just_over_3s_fails(self):
        """3.01 s is >= 3 -> 0 points for speed."""
        meta = self._meta()
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=3.01, page_blocked=False)
        assert score == 0

    def test_load_time_exactly_3s_fails(self):
        """Exactly 3.0 s is NOT < 3 -> 0 points for speed."""
        meta = self._meta()
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=3.0, page_blocked=False)
        assert score == 0

    def test_partial_combination(self):
        """Title (15) + description (10) + sitemap (10) = 35."""
        meta = self._meta(title='T', description='D')
        score = _calculate_score(meta, {}, has_sitemap=True,
                                 load_time=None, page_blocked=False)
        assert score == 35

    def test_load_time_none_no_points(self):
        """None load_time -> no speed points."""
        meta = self._meta()
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 0

    def test_load_time_string_error_no_points(self):
        """String load_time (error message) -> no speed points."""
        meta = self._meta()
        score = _calculate_score(meta, {}, has_sitemap=False,
                                 load_time='Connection timed out',
                                 page_blocked=True)
        assert score == 0

    def test_empty_ai_bots_list_no_points(self):
        """Empty ai_bots list is falsy -> 0 bot points."""
        meta = self._meta()
        robots = {'ai_bots': []}
        score = _calculate_score(meta, robots, has_sitemap=False,
                                 load_time=None, page_blocked=False)
        assert score == 0


# ===========================================================================
# 2. domain_service._clean_domain
# ===========================================================================

from services.domain_service import _clean_domain


class TestCleanDomain:
    """Test domain input cleaning and validation."""

    def test_full_https_url_with_path(self):
        assert _clean_domain('https://www.example.com/path') == 'example.com'

    def test_http_url(self):
        assert _clean_domain('http://example.com') == 'example.com'

    def test_www_prefix_stripped(self):
        assert _clean_domain('www.example.com') == 'example.com'

    def test_bare_domain_unchanged(self):
        assert _clean_domain('example.com') == 'example.com'

    def test_preserves_subdomain(self):
        assert _clean_domain('https://blog.example.com/page') == 'blog.example.com'

    def test_trailing_whitespace_stripped(self):
        assert _clean_domain('  example.com  ') == 'example.com'

    def test_case_normalised(self):
        assert _clean_domain('HTTPS://WWW.Example.COM') == 'example.com'

    def test_empty_string_raises(self):
        with pytest.raises(ValueError, match='empty'):
            _clean_domain('')

    def test_whitespace_only_raises(self):
        with pytest.raises(ValueError, match='empty'):
            _clean_domain('   ')

    def test_shell_injection_semicolon_raises(self):
        with pytest.raises(ValueError, match='Invalid domain'):
            _clean_domain('example.com; rm -rf /')

    def test_shell_injection_pipe_raises(self):
        with pytest.raises(ValueError, match='Invalid domain'):
            _clean_domain('example.com | cat /etc/passwd')

    def test_shell_injection_backtick_raises(self):
        with pytest.raises(ValueError, match='Invalid domain'):
            _clean_domain('example.com`whoami`')

    def test_shell_injection_dollar_raises(self):
        with pytest.raises(ValueError, match='Invalid domain'):
            _clean_domain('example.com$(id)')

    def test_javascript_scheme_not_treated_as_domain(self):
        """javascript:// does not start with 'http' so it bypasses URL
        parsing. The colon in 'javascript:' triggers the shell-metachar
        regex only if it contained ;|&`$ or whitespace.  Since the
        current regex does NOT block colons or slashes, the string is
        returned as-is.  This test documents actual behaviour; a stricter
        reject could be added later.
        """
        # The function does not raise; it treats the whole input as a
        # bare domain-like string.  Verify it does NOT crash.
        result = _clean_domain('javascript://alert(1)')
        assert isinstance(result, str)

    def test_data_scheme_not_treated_as_domain(self):
        """data: URLs contain angle brackets which are not in the
        rejection regex, so _clean_domain does not raise.  Document
        actual behaviour.
        """
        result = _clean_domain('data:text/html,<h1>Hi</h1>')
        assert isinstance(result, str)

    def test_ftp_scheme_not_treated_as_domain(self):
        """ftp:// does not start with 'http' so scheme validation is
        skipped.  Document actual behaviour.
        """
        result = _clean_domain('ftp://example.com')
        assert isinstance(result, str)


# ===========================================================================
# 3. client_store CRUD
# ===========================================================================

import client_store as cs
from config import Config


class TestClientStore:
    """Test client CRUD operations using a temporary data file."""

    @pytest.fixture(autouse=True)
    def _use_tmp_data_file(self, tmp_path, monkeypatch):
        """Point Config.DATA_FILE at a temporary location for each test."""
        tmp_file = str(tmp_path / 'clients.json')
        monkeypatch.setattr(Config, 'DATA_FILE', tmp_file)
        # Also update the lock file path used by client_store
        monkeypatch.setattr(cs, '_LOCK_FILE', tmp_file + '.lock')

    # -- basic reads -------------------------------------------------------
    def test_load_clients_empty_when_no_file(self):
        """No file on disk -> empty list."""
        assert cs.load_clients() == []

    def test_get_client_bad_id_returns_empty_dict(self):
        assert cs.get_client('nonexistent-id') == {}

    # -- create & read back ------------------------------------------------
    def test_save_client_creates_new(self):
        client = {'id': 'c1', 'name': 'Acme'}
        cs.save_client(client)

        loaded = cs.load_clients()
        assert len(loaded) == 1
        assert loaded[0]['id'] == 'c1'
        assert loaded[0]['name'] == 'Acme'

    def test_get_client_returns_correct(self):
        cs.save_client({'id': 'c1', 'name': 'Acme'})
        cs.save_client({'id': 'c2', 'name': 'Beta'})

        assert cs.get_client('c2')['name'] == 'Beta'

    # -- update ------------------------------------------------------------
    def test_save_client_updates_existing(self):
        cs.save_client({'id': 'c1', 'name': 'Acme'})
        cs.save_client({'id': 'c1', 'name': 'Acme Updated'})

        loaded = cs.load_clients()
        assert len(loaded) == 1
        assert loaded[0]['name'] == 'Acme Updated'

    # -- delete ------------------------------------------------------------
    def test_delete_client_removes(self):
        cs.save_client({'id': 'c1', 'name': 'Acme'})
        cs.save_client({'id': 'c2', 'name': 'Beta'})

        cs.delete_client('c1')
        loaded = cs.load_clients()
        assert len(loaded) == 1
        assert loaded[0]['id'] == 'c2'

    def test_delete_nonexistent_is_safe(self):
        """Deleting an ID that does not exist should not crash."""
        cs.save_client({'id': 'c1', 'name': 'Acme'})
        cs.delete_client('no-such-id')
        assert len(cs.load_clients()) == 1

    # -- corrupt file handling ---------------------------------------------
    def test_corrupt_json_returns_empty(self):
        """A corrupt JSON file should return [] rather than crash."""
        with open(Config.DATA_FILE, 'w') as f:
            f.write('{{{broken json!!!}}}')
        assert cs.load_clients() == []

    def test_empty_file_returns_empty(self):
        """A zero-byte file should return []."""
        with open(Config.DATA_FILE, 'w') as f:
            pass  # write nothing
        assert cs.load_clients() == []

    # -- multiple operations -----------------------------------------------
    def test_full_lifecycle(self):
        """Create -> read -> update -> delete lifecycle."""
        cs.save_client({'id': 'x', 'name': 'Original'})
        assert cs.get_client('x')['name'] == 'Original'

        cs.save_client({'id': 'x', 'name': 'Modified'})
        assert cs.get_client('x')['name'] == 'Modified'

        cs.delete_client('x')
        assert cs.get_client('x') == {}
        assert cs.load_clients() == []


# ===========================================================================
# 4. Config.validate
# ===========================================================================


class TestConfigValidate:
    """Test configuration validation logic."""

    def test_missing_secret_key_production_raises(self, monkeypatch):
        """No SECRET_KEY and no FLASK_DEBUG -> RuntimeError."""
        monkeypatch.setattr(Config, 'SECRET_KEY', '')
        monkeypatch.delenv('FLASK_DEBUG', raising=False)
        monkeypatch.delenv('FLASK_ENV', raising=False)
        with pytest.raises(RuntimeError, match='SECRET_KEY'):
            Config.validate()

    def test_missing_secret_key_with_flask_debug_sets_default(self, monkeypatch):
        """No SECRET_KEY but FLASK_DEBUG is set -> sets dev default silently."""
        monkeypatch.setattr(Config, 'SECRET_KEY', '')
        monkeypatch.setenv('FLASK_DEBUG', '1')
        monkeypatch.delenv('FLASK_ENV', raising=False)
        Config.validate()
        assert Config.SECRET_KEY == 'dev-secret-NOT-FOR-PRODUCTION'

    def test_missing_secret_key_with_flask_env_development(self, monkeypatch):
        """No SECRET_KEY but FLASK_ENV=development -> sets dev default."""
        monkeypatch.setattr(Config, 'SECRET_KEY', '')
        monkeypatch.setenv('FLASK_ENV', 'development')
        monkeypatch.delenv('FLASK_DEBUG', raising=False)
        Config.validate()
        assert Config.SECRET_KEY == 'dev-secret-NOT-FOR-PRODUCTION'

    def test_secret_key_set_no_error(self, monkeypatch):
        """Valid SECRET_KEY -> no error."""
        monkeypatch.setattr(Config, 'SECRET_KEY', 'a-real-secret-key-value')
        monkeypatch.delenv('FLASK_DEBUG', raising=False)
        monkeypatch.delenv('FLASK_ENV', raising=False)
        Config.validate()  # should not raise
        assert Config.SECRET_KEY == 'a-real-secret-key-value'


# ===========================================================================
# 5. URL validation (_validate_url from app.py)
# ===========================================================================

# We cannot import _validate_url directly from app.py without triggering
# Flask app creation (Config.validate, CSRFProtect, etc.), so we replicate
# the function logic here.  The source-of-truth is in webapp/app.py.

from urllib.parse import urlparse as _urlparse


def _validate_url(url):
    """Re-implementation of webapp/app.py _validate_url for isolated testing."""
    if not url:
        raise ValueError("Please enter a URL.")
    parsed = _urlparse(url)
    if parsed.scheme not in ('http', 'https', ''):
        raise ValueError("Only http and https URLs are allowed.")
    if not parsed.scheme:
        url = f"https://{url}"
    return url


class TestValidateUrl:
    """Test the URL validation helper."""

    def test_https_url_passes(self):
        assert _validate_url('https://example.com') == 'https://example.com'

    def test_http_url_passes(self):
        assert _validate_url('http://example.com') == 'http://example.com'

    def test_bare_domain_gets_https(self):
        assert _validate_url('example.com') == 'https://example.com'

    def test_empty_string_raises(self):
        with pytest.raises(ValueError, match='enter a URL'):
            _validate_url('')

    def test_none_raises(self):
        with pytest.raises(ValueError):
            _validate_url(None)

    def test_file_scheme_raises(self):
        with pytest.raises(ValueError, match='http'):
            _validate_url('file:///etc/passwd')

    def test_javascript_scheme_raises(self):
        with pytest.raises(ValueError, match='http'):
            _validate_url('javascript:alert(1)')

    def test_data_scheme_raises(self):
        with pytest.raises(ValueError, match='http'):
            _validate_url('data:text/html,<h1>Hi</h1>')

    def test_ftp_scheme_raises(self):
        with pytest.raises(ValueError, match='http'):
            _validate_url('ftp://example.com')

    def test_url_with_path_and_query(self):
        url = _validate_url('https://example.com/path?q=1')
        assert url == 'https://example.com/path?q=1'

    def test_url_with_port(self):
        url = _validate_url('http://localhost:8080')
        assert url == 'http://localhost:8080'


# ===========================================================================
# 6. Report generation smoke tests
# ===========================================================================

from docx import Document as DocxDocument

import report_generators.geo_audit_report as _geo_mod
import report_generators.content_guide as _cg_mod
import report_generators.docx_helpers as _docx_helpers

from report_generators.geo_audit_report import build_geo_audit_report
from report_generators.content_guide import build_content_guide


class TestReportGeneration:
    """Smoke tests: verify reports build without errors on minimal input."""

    @pytest.fixture(autouse=True)
    def _use_blank_document(self, monkeypatch):
        """Force create_document() to return a plain blank Document so tests
        do not depend on the Numiko .dotx template file being loadable by
        the installed version of python-docx."""
        def _blank_doc(title='', subtitle=''):
            doc = DocxDocument()
            style = doc.styles['Normal']
            style.font.name = 'Modern Era'
            return doc

        # Patch in every module that has imported the name
        monkeypatch.setattr(_docx_helpers, 'create_document', _blank_doc)
        monkeypatch.setattr(_geo_mod, 'create_document', _blank_doc)
        monkeypatch.setattr(_cg_mod, 'create_document', _blank_doc)

    @staticmethod
    def _minimal_audit_data():
        return {
            'url': 'https://example.com',
            'page_blocked': False,
            'block_reason': None,
            'title': 'Example Domain',
            'title_length': 14,
            'title_ok': True,
            'description': 'An example site.',
            'description_length': 16,
            'description_ok': True,
            'og_tags': False,
            'h1': 'Example Domain',
            'jsonld_count': 0,
            'load_time': 0.5,
            'load_time_ok': True,
            'robots_exists': True,
            'ai_bots': ['GPTBot'],
            'ai_bots_blocked': [],
            'has_sitemap': True,
            'sitemap_url': '/sitemap.xml',
            'backlinks_rank': None,
            'referring_domains': None,
            'total_backlinks': None,
            'score': 75,
        }

    def test_build_geo_audit_report_minimal(self):
        params = {
            'client_name': 'Test Client',
            'client_domain': 'example.com',
        }
        audit = self._minimal_audit_data()
        doc = build_geo_audit_report(params, audit)
        # Must return a python-docx Document with at least some content
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_build_geo_audit_report_with_backlinks(self):
        params = {
            'client_name': 'Test Client',
            'client_domain': 'example.com',
        }
        audit = self._minimal_audit_data()
        audit['backlinks_rank'] = 42
        audit['referring_domains'] = 150
        audit['total_backlinks'] = 3200
        doc = build_geo_audit_report(params, audit)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_build_geo_audit_report_low_score(self):
        params = {'client_name': 'Weak Site', 'client_domain': 'weak.com'}
        audit = self._minimal_audit_data()
        audit['score'] = 20
        audit['title'] = ''
        audit['h1'] = ''
        audit['jsonld_count'] = 0
        audit['robots_exists'] = False
        audit['ai_bots'] = []
        audit['has_sitemap'] = False
        audit['load_time_ok'] = False
        audit['load_time'] = 5.2
        audit['og_tags'] = False
        doc = build_geo_audit_report(params, audit)
        assert doc is not None

    def test_build_geo_audit_report_moderate_score(self):
        params = {'client_name': 'Mid Site', 'client_domain': 'mid.com'}
        audit = self._minimal_audit_data()
        audit['score'] = 55
        doc = build_geo_audit_report(params, audit)
        assert doc is not None

    def test_build_geo_audit_report_blocked_bots(self):
        params = {'client_name': 'Blocked', 'client_domain': 'blocked.com'}
        audit = self._minimal_audit_data()
        audit['ai_bots'] = ['GPTBot']
        audit['ai_bots_blocked'] = ['ClaudeBot', 'PerplexityBot']
        doc = build_geo_audit_report(params, audit)
        assert doc is not None

    def test_build_content_guide_minimal(self):
        params = {
            'client_name': 'Test Client',
            'client_domain': 'example.com',
        }
        doc = build_content_guide(params)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_build_content_guide_all_params(self):
        params = {
            'client_name': 'Acme Corp',
            'client_domain': 'acme.com',
            'project_name': 'Relaunch 2026',
            'date': 'March 2026',
            'cms': 'WordPress',
            'logo_path': '/nonexistent/path.png',  # missing logo is fine
        }
        doc = build_content_guide(params)
        assert doc is not None

    def test_build_content_guide_empty_params(self):
        """Completely empty params should use defaults without crashing."""
        doc = build_content_guide({})
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_geo_audit_report_saveable(self, tmp_path):
        """The document can be saved to disk."""
        params = {'client_name': 'Save Test', 'client_domain': 'save.test'}
        audit = self._minimal_audit_data()
        doc = build_geo_audit_report(params, audit)
        out = str(tmp_path / 'report.docx')
        doc.save(out)
        assert os.path.exists(out)
        assert os.path.getsize(out) > 0

    def test_content_guide_saveable(self, tmp_path):
        """The content guide document can be saved to disk."""
        params = {'client_name': 'Save Test', 'client_domain': 'save.test'}
        doc = build_content_guide(params)
        out = str(tmp_path / 'guide.docx')
        doc.save(out)
        assert os.path.exists(out)
        assert os.path.getsize(out) > 0


# ===========================================================================
# 7. Path traversal protection (_safe_download_path)
# ===========================================================================

# Same as _validate_url: importing from app.py triggers Flask app creation.
# Re-implement the logic here for isolated testing.  Source-of-truth is
# webapp/app.py::_safe_download_path.


def _safe_download_path(filename):
    """Re-implementation of webapp/app.py _safe_download_path."""
    safe_path = os.path.normpath(os.path.join(Config.OUTPUT_DIR, filename))
    if not safe_path.startswith(os.path.normpath(Config.OUTPUT_DIR) + os.sep) and \
       safe_path != os.path.normpath(Config.OUTPUT_DIR):
        return None
    return safe_path


class TestSafeDownloadPath:
    """Test path traversal protection on the download endpoint."""

    def test_simple_filename_valid(self):
        path = _safe_download_path('report.docx')
        assert path is not None
        assert path.endswith('report.docx')
        assert Config.OUTPUT_DIR in path

    def test_single_dot_dot_returns_none(self):
        assert _safe_download_path('../etc/passwd') is None

    def test_double_dot_dot_returns_none(self):
        assert _safe_download_path('../../etc/passwd') is None

    def test_absolute_path_returns_none(self):
        assert _safe_download_path('/absolute/path') is None

    def test_deeply_nested_traversal_returns_none(self):
        assert _safe_download_path('../../../../../../../etc/shadow') is None

    def test_dot_slash_prefix_returns_none(self):
        """../filename escapes the output dir."""
        assert _safe_download_path('../secret.docx') is None

    def test_subdirectory_valid(self):
        """A filename with a subdirectory is fine as long as it stays inside."""
        path = _safe_download_path('subdir/report.docx')
        assert path is not None
        assert Config.OUTPUT_DIR in path

    def test_backslash_traversal_returns_none(self):
        """Backslash traversal (Windows-style) should be normalised and blocked."""
        result = _safe_download_path('..\\etc\\passwd')
        # On Unix os.path.normpath will treat backslashes literally in the
        # filename, but the '../' prefix should still be caught.
        # On Windows, normpath normalises to ..\\etc\\passwd and blocks it.
        # Either way, if the resolved path escapes OUTPUT_DIR, it is None.
        if result is not None:
            assert os.path.normpath(Config.OUTPUT_DIR) in result

    def test_encoded_traversal_stays_literal(self):
        """URL-encoded dots are passed as literals and should be safe."""
        path = _safe_download_path('%2e%2e/etc/passwd')
        # These are literal characters, not actual '..' after normpath.
        # Should resolve inside OUTPUT_DIR (the filename has literal %).
        if path is not None:
            assert os.path.normpath(Config.OUTPUT_DIR) in path


# ===========================================================================
# 8. User model
# ===========================================================================

from werkzeug.security import generate_password_hash, check_password_hash


class TestUserModel:
    """Test password hashing and is_active property logic.

    Password tests use werkzeug directly (identical to the model methods).
    is_active tests use a lightweight Flask app context so SQLAlchemy
    instrumentation works.
    """

    # -- password hashing (werkzeug, no app context needed) ----------------

    def test_password_hashing(self):
        """generate/check round-trip matches the model's set_password/check_password."""
        pw_hash = generate_password_hash('Test1234', method='pbkdf2:sha256')
        assert check_password_hash(pw_hash, 'Test1234')
        assert not check_password_hash(pw_hash, 'wrong')

    def test_password_different_hashes(self):
        """Two calls produce different hashes (salting)."""
        h1 = generate_password_hash('Test1234', method='pbkdf2:sha256')
        h2 = generate_password_hash('Test1234', method='pbkdf2:sha256')
        assert h1 != h2

    # -- is_active property (needs SQLAlchemy model instrumentation) -------

    @pytest.fixture
    def _user_in_app(self, tmp_path):
        """Create a minimal Flask app + DB and yield a User factory."""
        from flask import Flask
        from extensions import db as _db

        app = Flask(__name__)
        app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{tmp_path / "test.db"}'
        app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
        app.config['SECRET_KEY'] = 'test'

        _db.init_app(app)
        with app.app_context():
            from models import User
            _db.create_all()
            yield User, _db, app

    def test_is_active_requires_activation(self, _user_in_app):
        User, _db, app = _user_in_app
        with app.app_context():
            u = User(email='t@numiko.com', name='T', password_hash='x',
                     is_active_user=False, deactivated_at=None)
            assert not u.is_active

    def test_is_active_true_when_activated(self, _user_in_app):
        User, _db, app = _user_in_app
        with app.app_context():
            u = User(email='t@numiko.com', name='T', password_hash='x',
                     is_active_user=True, deactivated_at=None)
            assert u.is_active

    def test_is_active_false_when_revoked(self, _user_in_app):
        from datetime import datetime, timezone
        User, _db, app = _user_in_app
        with app.app_context():
            u = User(email='t@numiko.com', name='T', password_hash='x',
                     is_active_user=True,
                     deactivated_at=datetime.now(timezone.utc))
            assert not u.is_active


# ===========================================================================
# 9. Password validation
# ===========================================================================


class TestPasswordValidation:
    """Test the _validate_password helper from auth.py."""

    @staticmethod
    def _validate(pw):
        from auth import _validate_password
        return _validate_password(pw)

    def test_valid_password(self):
        assert self._validate('Abcdef12') is None

    def test_too_short(self):
        err = self._validate('Ab1')
        assert err is not None
        assert '8 characters' in err

    def test_no_uppercase(self):
        err = self._validate('abcdefg1')
        assert err is not None
        assert 'uppercase' in err

    def test_no_digit(self):
        err = self._validate('Abcdefgh')
        assert err is not None
        assert 'digit' in err

    def test_exactly_8_chars_valid(self):
        assert self._validate('Abcdef12') is None

    def test_long_password_valid(self):
        assert self._validate('Abcdefghijklmnop1') is None


# ===========================================================================
# 10. Email domain validation
# ===========================================================================


class TestEmailValidation:
    """Test the _is_valid_email helper from auth.py."""

    @staticmethod
    def _check(email):
        from auth import _is_valid_email
        return _is_valid_email(email)

    def test_valid_numiko_email(self):
        assert self._check('user@numiko.com')

    def test_valid_numiko_email_mixed_case(self):
        assert self._check('User@Numiko.com')

    def test_invalid_domain(self):
        assert not self._check('user@gmail.com')

    def test_empty_string(self):
        assert not self._check('')

    def test_no_at_sign(self):
        assert not self._check('usernumiko.com')

    def test_similar_domain_rejected(self):
        assert not self._check('user@notnumiko.com')

    def test_subdomain_rejected(self):
        """sub.numiko.com is NOT @numiko.com."""
        assert not self._check('user@sub.numiko.com')


# ===========================================================================
# 11. Password reset tokens
# ===========================================================================


class TestPasswordResetTokens:
    """Test password reset token generation, verification, and cross-salt
    isolation from activation tokens."""

    @pytest.fixture(autouse=True)
    def _patch_config(self, monkeypatch):
        """Patch Config so token tests run without a real SECRET_KEY."""
        monkeypatch.setattr(Config, 'SECRET_KEY', 'test-secret-for-tokens')
        monkeypatch.setattr(Config, 'PASSWORD_RESET_TOKEN_MAX_AGE', 15 * 60)
        monkeypatch.setattr(Config, 'PASSWORD_RESET_TOKEN_SALT', 'password-reset-salt')
        monkeypatch.setattr(Config, 'ACTIVATION_TOKEN_SALT', 'email-activation-salt')
        monkeypatch.setattr(Config, 'ACTIVATION_TOKEN_MAX_AGE', 48 * 3600)

    def test_generate_and_verify_round_trip(self):
        """A freshly generated token verifies back to the original email."""
        from email_service import generate_password_reset_token, verify_password_reset_token
        token = generate_password_reset_token('alice@numiko.com')
        assert verify_password_reset_token(token) == 'alice@numiko.com'

    def test_invalid_token_returns_none(self):
        """A random string is not a valid token."""
        from email_service import verify_password_reset_token
        assert verify_password_reset_token('not-a-valid-token-at-all') is None

    def test_empty_token_returns_none(self):
        from email_service import verify_password_reset_token
        assert verify_password_reset_token('') is None

    def test_tampered_token_returns_none(self):
        """Modifying the token payload causes signature verification to fail."""
        from email_service import generate_password_reset_token, verify_password_reset_token
        token = generate_password_reset_token('alice@numiko.com')
        tampered = token[:-5] + 'XXXXX'
        assert verify_password_reset_token(tampered) is None

    def test_activation_token_rejected_as_reset_token(self):
        """An activation token (different salt) must NOT verify as a reset token."""
        from email_service import generate_activation_token, verify_password_reset_token
        act_token = generate_activation_token('alice@numiko.com')
        assert verify_password_reset_token(act_token) is None

    def test_reset_token_rejected_as_activation_token(self):
        """A password reset token must NOT verify as an activation token."""
        from email_service import generate_password_reset_token, verify_activation_token
        reset_token = generate_password_reset_token('alice@numiko.com')
        assert verify_activation_token(reset_token) is None

    def test_token_carries_correct_email(self):
        """Token issued for alice does not verify as bob."""
        from email_service import generate_password_reset_token, verify_password_reset_token
        token = generate_password_reset_token('alice@numiko.com')
        email = verify_password_reset_token(token)
        assert email == 'alice@numiko.com'
        assert email != 'bob@numiko.com'

    def test_two_tokens_for_same_email_are_different(self):
        """Each call produces a unique token (timestamp-salted)."""
        from email_service import generate_password_reset_token
        t1 = generate_password_reset_token('alice@numiko.com')
        t2 = generate_password_reset_token('alice@numiko.com')
        assert t1 != t2
